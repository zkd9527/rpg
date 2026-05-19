#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown文件转换为Word格式的脚本
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown import markdown
from html.parser import HTMLParser

class MDToWord:
    def __init__(self):
        self.doc = Document()
        self.current_list_level = 0
        
    def convert_file(self, md_file, output_file):
        """将Markdown文件转换为Word文件"""
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        self.convert_markdown(content)
        self.doc.save(output_file)
        print(f"✓ 已生成：{output_file}")
    
    def convert_markdown(self, content):
        """处理Markdown内容"""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # 跳过空行
            if not line.strip():
                i += 1
                continue
            
            # 处理标题
            if line.startswith('# '):
                self.add_title1(line[2:].strip())
            elif line.startswith('## '):
                self.add_title2(line[3:].strip())
            elif line.startswith('### '):
                self.add_title3(line[4:].strip())
            elif line.startswith('#### '):
                self.add_title4(line[5:].strip())
            
            # 处理列表
            elif line.startswith('- '):
                self.add_bullet_point(line[2:].strip())
            elif re.match(r'^\d+\. ', line):
                match = re.match(r'^\d+\. (.+)', line)
                if match:
                    self.add_numbered_point(match.group(1).strip())
            
            # 处理代码块
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                self.add_code_block('\n'.join(code_lines))
                i += 1
                continue
            
            # 处理引用
            elif line.startswith('> '):
                self.add_quote(line[2:].strip())
            
            # 处理表格
            elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                table_lines = [line]
                i += 1
                table_lines.append(lines[i])  # 添加分隔符行
                i += 1
                # 读取表格的其他行
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self.add_table(table_lines)
                i -= 1  # 回退一行因为循环会增加i
            
            # 处理普通文本
            else:
                # 检查是否是列表项（忽略缩进）
                if not any(line.lstrip().startswith(x) for x in ['- ', '* ', '> ']):
                    self.add_paragraph(line)
            
            i += 1
    
    def add_title1(self, text):
        """添加一级标题"""
        p = self.doc.add_heading(text, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def add_title2(self, text):
        """添加二级标题"""
        p = self.doc.add_heading(text, level=2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def add_title3(self, text):
        """添加三级标题"""
        p = self.doc.add_heading(text, level=3)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def add_title4(self, text):
        """添加四级标题"""
        p = self.doc.add_heading(text, level=4)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def add_paragraph(self, text):
        """添加段落"""
        if text.strip():
            # 处理内联格式
            p = self.doc.add_paragraph()
            self.add_formatted_text(p, text)
    
    def add_formatted_text(self, paragraph, text):
        """添加格式化文本（支持粗体、斜体、代码等）"""
        # 简单的格式化处理
        while text:
            # 处理粗体
            match = re.search(r'\*\*(.+?)\*\*', text)
            if match:
                before = text[:match.start()]
                if before:
                    paragraph.add_run(before)
                run = paragraph.add_run(match.group(1))
                run.bold = True
                text = text[match.end():]
                continue
            
            # 处理斜体
            match = re.search(r'\*(.+?)\*', text)
            if match:
                before = text[:match.start()]
                if before:
                    paragraph.add_run(before)
                run = paragraph.add_run(match.group(1))
                run.italic = True
                text = text[match.end():]
                continue
            
            # 处理行内代码
            match = re.search(r'`(.+?)`', text)
            if match:
                before = text[:match.start()]
                if before:
                    paragraph.add_run(before)
                run = paragraph.add_run(match.group(1))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                text = text[match.end():]
                continue
            
            # 处理链接
            match = re.search(r'\[(.+?)\]\((.+?)\)', text)
            if match:
                before = text[:match.start()]
                if before:
                    paragraph.add_run(before)
                run = paragraph.add_run(match.group(1))
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
                text = text[match.end():]
                continue
            
            # 没有特殊格式，添加整个文本
            if text:
                paragraph.add_run(text)
            break
    
    def add_bullet_point(self, text):
        """添加项目符号"""
        p = self.doc.add_paragraph(text, style='List Bullet')
    
    def add_numbered_point(self, text):
        """添加编号列表"""
        p = self.doc.add_paragraph(text, style='List Number')
    
    def add_quote(self, text):
        """添加引用"""
        p = self.doc.add_paragraph(text, style='Normal')
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        for run in p.runs:
            run.italic = True
    
    def add_code_block(self, code):
        """添加代码块"""
        p = self.doc.add_paragraph(code, style='Normal')
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        # 设置灰色背景
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F0F0F0')
        p._element.get_or_add_pPr().append(shading_elm)
    
    def add_table(self, table_lines):
        """添加表格"""
        if len(table_lines) < 3:
            return
        
        # 解析表格行
        rows = []
        for line in table_lines:
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                rows.append(cells)
        
        if not rows:
            return
        
        # 创建表格
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'
        
        # 填充表格内容
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text

def main():
    """主函数"""
    print("开始转换Markdown文件为Word格式...\n")
    
    # 转换的文件列表
    files_to_convert = [
        ("论文第四章优化版.md", "D:\\论文优化文档\\论文第四章优化版.docx"),
        ("优化说明.md", "D:\\论文优化文档\\优化说明.docx"),
    ]
    
    for md_file, output_file in files_to_convert:
        if os.path.exists(md_file):
            try:
                converter = MDToWord()
                converter.convert_file(md_file, output_file)
            except Exception as e:
                print(f"✗ 转换 {md_file} 失败: {e}")
        else:
            print(f"✗ 文件不存在: {md_file}")
    
    print("\n转换完成！")

if __name__ == "__main__":
    main()
